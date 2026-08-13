# -*- coding: utf-8 -*-
"""生成季度新番信息 Word 文档（含封面、每部番封面图与美化排版）。"""
import datetime
import os
import sys

from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from common import (LIGHT_BG, add_cover, add_footer_pagenum, add_heading, add_hr, add_para,
                    cell_text, colorize_score_cell, download_cover, init_doc, make_table,
                    season_label, set_run, setup_logging, shade_cell, spacer, summary_cn)


def _fmt_rating(item):
    r = item.get("rating")
    if r is None:
        return "暂无评分"
    s = "%.1f" % r
    total = item.get("rating_total") or 0
    rank = item.get("rank")
    out = "%s（%d 人评分" % (s, total)
    if rank:
        out += "，排名 #%d" % rank
    out += "）"
    return out


def _fmt_time(item):
    parts = [item.get("weekday", "")]
    if item.get("time"):
        parts.append(item["time"])
    txt = " ".join(p for p in parts if p)
    if item.get("eps"):
        txt += "　" + str(item["eps"])
    return txt or "—"


def build(yuc_data, bgm_items, year, start_month, out_path, generated_at=None):
    generated_at = generated_at or datetime.date.today().isoformat()
    label = season_label(year, start_month)
    items = list(bgm_items)
    total = len(items)
    matched = sum(1 for x in items if x.get("bgm_id"))
    unmatched = total - matched
    originals = sum(1 for x in items if x.get("original") is True)
    adapted = sum(1 for x in items if x.get("original") is False)
    weekday_order = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
    wd_count = {}
    for x in items:
        wd = x.get("weekday") or "未知"
        wd_count[wd] = wd_count.get(wd, 0) + 1

    doc = init_doc()
    doc.core_properties.title = "%s新番信息" % label
    doc.core_properties.author = "Anime Auto"
    add_cover(
        doc,
        "%s新番信息" % label,
        ["日本 TV 动画 · 季度新番全收录", "数据来源：yuc.wiki（番剧清单） + Bangumi（详细信息与评分）"],
        date_line="生成日期：%s" % generated_at,
    )
    add_footer_pagenum(doc)

    add_heading(doc, "一、统计概览", level=1)
    stats = [
        ["收录作品", "%d 部" % total],
        ["已匹配 Bangumi 条目", "%d 部" % matched],
        ["未匹配（仅保留 yuc.wiki 标题）", "%d 部" % unmatched],
        ["原创动画", "%d 部" % originals],
        ["改编动画", "%d 部" % adapted],
    ]
    for wd in weekday_order:
        if wd_count.get(wd):
            stats.append(["放送于 %s" % wd, "%d 部" % wd_count[wd]])
    make_table(doc, ["项目", "数量"], stats, widths_cm=[6.0, 9.5], font_size=10)
    spacer(doc)

    add_heading(doc, "二、新番一览表", level=1)
    rows = []
    for i, x in enumerate(items, start=1):
        otype = {True: "原创", False: "改编", None: "—"}.get(x.get("original"), "—")
        rows.append([
            str(i),
            x.get("title", ""),
            x.get("weekday", ""),
            x.get("time", ""),
            x.get("eps", ""),
            otype,
            x.get("studio", "") or "—",
            "%.1f" % x["rating"] if x.get("rating") is not None else "—",
        ])
    t = make_table(doc,
                   ["序", "标题", "星期", "时间", "话数", "原创", "制作公司", "评分"],
                   rows,
                   widths_cm=[0.9, 4.6, 1.4, 1.6, 1.4, 1.3, 3.4, 1.5],
                   font_size=8)
    for row in t.rows[1:]:
        colorize_score_cell(row.cells[7])
    spacer(doc)

    add_heading(doc, "三、每部番详细信息", level=1)
    logger = setup_logging()
    for i, x in enumerate(items, start=1):
        add_heading(doc, "%d. %s" % (i, x.get("title", "")), level=2)
        fields = [
            ["放送", _fmt_time(x)],
            ["是否原创", {True: "原创", False: "改编（原作：%s）" % (x.get("source") or "—"), None: "未知"}.get(x.get("original"), "—")],
            ["原作", x.get("source") or "—"],
            ["导演", x.get("director") or "—"],
            ["系列构成 / 脚本", x.get("script") or "—"],
            ["动画制作", x.get("studio") or "—"],
            ["主要声优", x.get("cast") or "—"],
            ["标签", x.get("tags") or "—"],
            ["Bangumi 评分", _fmt_rating(x)],
            ["Bangumi 链接", x.get("bgm_url") or "未找到对应条目"],
        ]
        cover = download_cover(x.get("bgm_id"), x.get("image_url"), logger=logger)
        if cover:
            outer = doc.add_table(rows=1, cols=2)
            outer.style = "Table Grid"
            outer.autofit = False
            left, right = outer.rows[0].cells
            left.width = Cm(3.8)
            right.width = Cm(12.2)
            left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            try:
                p = left.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(cover, width=Cm(3.2))
            except Exception as e:
                logger.warning("封面插入失败「%s」：%s", x.get("title"), e)
                outer._element.getparent().remove(outer._element)
                cover = None
        if cover:
            inner = right.add_table(rows=len(fields), cols=2)
            inner.style = "Table Grid"
            inner.autofit = False
            for ri, (k, v) in enumerate(fields):
                c0 = inner.rows[ri].cells[0]
                c1 = inner.rows[ri].cells[1]
                cell_text(c0, k, size=9, bold=True)
                shade_cell(c0, LIGHT_BG)
                cell_text(c1, v, size=9)
                c0.width = Cm(2.7)
                c1.width = Cm(9.2)
        else:
            t2 = make_table(doc, ["项目", "内容"], fields, widths_cm=[3.0, 13.0], font_size=9)
            for row in t2.rows:
                shade_cell(row.cells[0], LIGHT_BG)
        spacer(doc, 2)
        add_para(doc, "简介：", size=9, bold=True, color="2B6CB0", space_after=2)
        summary = summary_cn(x)
        add_para(doc, summary, size=9, space_after=10)
        add_hr(doc, color="CBD5E1", size=8, space_after=10)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)

    # 同步生成同名 JSON（与 docx 放在同一季度文件夹），供「番剧记录」应用直接导入
    try:
        import json
        json_path = os.path.splitext(out_path)[0] + ".json"
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(
                {"season": "%d%02d" % (year, start_month), "generated_at": generated_at, "items": items},
                f, ensure_ascii=False, indent=1,
            )
    except Exception as e:
        logger.warning("JSON 生成失败：%s", e)

    return out_path