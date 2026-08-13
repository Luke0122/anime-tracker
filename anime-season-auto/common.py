# -*- coding: utf-8 -*-
"""共享工具：路径、配置、日志、JSON、日期计算、docx 样式助手。"""
import datetime
import json
import logging
import os
import re
import sys

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_ROOT = os.path.dirname(BASE_DIR)          # D:\ANIME\日本TV动画信息
RUNTIME_DIR = os.path.join(BASE_DIR, "_runtime")
LOG_DIR = os.path.join(RUNTIME_DIR, "logs")
CACHE_DIR = os.path.join(RUNTIME_DIR, "data")
COVER_DIR = os.path.join(RUNTIME_DIR, "covers")
STATE_FILE = os.path.join(RUNTIME_DIR, "state.json")
CONFIG_FILE = os.path.join(BASE_DIR, "config.json")

SEASON_LABEL = {1: "1月", 4: "4月", 7: "7月", 10: "10月"}
FONT = "微软雅黑"
ACCENT = "2B6CB0"
HEADER_BG = "2B6CB0"
LIGHT_BG = "EAF2FB"
GRAY_BG = "F2F2F2"


def ensure_dirs():
    for d in (RUNTIME_DIR, LOG_DIR, CACHE_DIR, COVER_DIR):
        os.makedirs(d, exist_ok=True)


def setup_logging():
    ensure_dirs()
    logger = logging.getLogger("anime")
    if logger.handlers:
        return logger
    logger.setLevel(logging.INFO)
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    fh = logging.FileHandler(
        os.path.join(LOG_DIR, datetime.date.today().strftime("%Y%m%d") + ".log"),
        encoding="utf-8",
    )
    fh.setFormatter(fmt)
    logger.addHandler(fh)
    try:
        sh = logging.StreamHandler(sys.stdout)
        sh.setFormatter(fmt)
        logger.addHandler(sh)
    except Exception:
        pass
    return logger


def reconfigure_console():
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass


def load_json(path, default=None):
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return default


def save_json(path, data):
    ensure_dirs()
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    os.replace(tmp, path)


def load_config():
    cfg = load_json(CONFIG_FILE, {}) or {}
    cfg.setdefault("base_dir", DATA_ROOT)
    cfg.setdefault("mail_to", "your-email@example.com")
    cfg.setdefault("smtp_host", "smtp.qq.com")
    cfg.setdefault("smtp_port", 465)
    cfg.setdefault("smtp_user", "your-email@example.com")
    cfg.setdefault("smtp_auth_code", "")
    cfg.setdefault("request_delay", 1.2)
    cfg.setdefault("max_voice_actors", 8)
    cfg.setdefault("top_tags", 10)
    cfg.setdefault("send_email", True)
    return cfg


def today():
    return datetime.date.today()


def season_key(year, start_month):
    return "%04d-%02d" % (year, start_month)


def season_label(year, start_month):
    return "%d年%s" % (year, SEASON_LABEL[start_month])


def season_dir(year, start_month):
    return os.path.join(DATA_ROOT, season_label(year, start_month))


def yuc_month(year, start_month):
    return "%04d%02d" % (year, start_month)


def add_months(year, month, delta):
    m = month - 1 + delta
    y = year + m // 12
    m = m % 12 + 1
    return y, m


def _has_kana(s):
    return any(
        "\u3041" <= ch <= "\u3096"          # 平假名
        or "\u30a1" <= ch <= "\u30fa"       # 片假名（不含・）
        or ch == "\u30fc"                     # 长音记号
        or "\uff66" <= ch <= "\uff9f"        # 半角片假名
        for ch in s
    )


_JP_VARIANT_TO_CN = {
    "伝": "传", "髙": "高", "黒": "黑", "縁": "缘", "桜": "樱", "渋": "涩",
    "歳": "岁", "髪": "发", "竜": "龙", "亜": "亚", "斎": "斋", "斉": "齐",
    "沢": "泽", "浜": "滨", "嶋": "岛", "焔": "焰", "呉": "吴", "埜": "野",
    "邉": "边", "姫": "姬", "嬢": "娘", "頬": "颊", "壊": "坏", "継": "继",
    "縦": "纵", "繋": "系", "巻": "卷", "圧": "压", "拡": "扩", "単": "单",
    "変": "变", "応": "应", "営": "营", "撃": "击", "検": "检", "権": "权",
    "済": "济", "渉": "涉", "経": "经", "総": "总", "続": "续", "覧": "览",
    "譲": "让", "豊": "丰", "賛": "赞", "隣": "邻", "顕": "显", "駆": "驱",
    "仮": "假", "価": "价", "勧": "劝", "収": "收", "図": "图", "囲": "围",
    "団": "团", "塩": "盐", "増": "增", "挙": "举", "摂": "摄", "暦": "历",
    "歴": "历", "気": "气", "涙": "泪", "焼": "烧", "獣": "兽", "畳": "叠",
    "穂": "穗", "窓": "窗", "絵": "绘", "縄": "绳", "聴": "听", "栄": "荣",
    "薬": "药", "覚": "觉", "観": "观", "訳": "译", "読": "读", "逓": "递",
    "郷": "乡", "醸": "酿", "関": "关", "闘": "斗", "険": "险", "隠": "隐",
    "髄": "髓", "鶏": "鸡", "黙": "默", "帰": "归", "剣": "剑", "戦": "战",
    "機": "机", "仏": "佛", "乗": "乘", "児": "儿", "冨": "富", "辺": "边",
    "﨑": "崎", "神": "神", "都": "都", "栃": "枥",
}
_JP_VARIANT_TRANS = str.maketrans(_JP_VARIANT_TO_CN)
_opencc = None


def _t2s(text):
    global _opencc
    try:
        if _opencc is None:
            from opencc import OpenCC
            _opencc = OpenCC("t2s")
        return _opencc.convert(text)
    except Exception:
        return text


def strip_kana(text):
    """去掉假名与长音记号，保留中文标点（含分隔点・）。"""
    out = []
    for ch in str(text or ""):
        o = ord(ch)
        if (0x3041 <= o <= 0x30FA or 0x30FC <= o <= 0x30FF
                or 0x31F0 <= o <= 0x31FF or 0xFF66 <= o <= 0xFF9F):
            continue
        out.append(ch)
    return "".join(out).strip()


def simplify_cn(text, remove_kana=False):
    """繁体/日文变体 -> 简体中文；可选去掉残留假名。"""
    if text is None:
        return ""
    s = str(text)
    s = _t2s(s)
    if _JP_VARIANT_TRANS:
        s = s.translate(_JP_VARIANT_TRANS)
    if remove_kana:
        s = strip_kana(s)
    return s


def summary_cn(item, fallback="（暂无中文简介，详情见 Bangumi 条目）"):
    """从 Bangumi 简介中提取中文部分（去掉日文原文），并转为简体中文。"""
    s = str(item.get("summary") or "")
    s = s.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not s:
        return "（暂无简介）"
    if "[简介原文]" in s:
        pre = s.split("[简介原文]")[0].strip()
        pre = "\n".join([x for x in pre.split("\n") if x.strip()])
        return simplify_cn(pre, remove_kana=True) if pre else fallback
    paras = [pp.strip() for pp in s.split("\n")]
    paras = [pp for pp in paras if pp]
    keep = []
    for pp in paras:
        if _has_kana(pp):
            break
        keep.append(pp)
    if keep:
        return simplify_cn("\n".join(keep), remove_kana=True)
    return fallback


def norm_title(s):
    s = (s or "").replace("　", "")
    return re.sub(r"[\s・·、,，！!？?：:（）()。．～—-]+", "", s)


# ---------------- docx 样式助手 ----------------

def init_doc():
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = None
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)
    for sec in doc.sections:
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.6)
        sec.right_margin = Cm(1.6)
    return doc


def set_run(run, size=9, bold=False, color=None, italic=False):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    rfonts.set(qn("w:eastAsia"), FONT)


def shade_cell(cell, color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcpr.append(shd)


def cell_text(cell, text, size=9, bold=False, color=None, align="left"):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=color)


def add_para(doc, text="", size=9, bold=False, color=None, align=None,
             space_after=6, space_before=0):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    set_run(run, size=size, bold=bold, color=color)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_heading(doc, text, level=1):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    if level == 0:
        run = p.add_run(text)
        set_run(run, size=22, bold=True, color=ACCENT)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
    elif level == 1:
        run = p.add_run(text)
        set_run(run, size=15, bold=True, color=ACCENT)
        p.paragraph_format.space_before = Pt(14)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), ACCENT)
        pbdr.append(bottom)
        pPr.append(pbdr)
    elif level == 2:
        run = p.add_run(text)
        set_run(run, size=12, bold=True, color="333333")
        p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    return p


def make_table(doc, headers, rows, widths_cm=None, font_size=9,
               header_color=HEADER_BG, first_col_bold=False, zebra=True,
               repeat_header=True):
    from docx.shared import Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell_text(cell, h, size=font_size, bold=True, color="FFFFFF", align="center")
        shade_cell(cell, header_color)
    if repeat_header:
        trPr = t.rows[0]._tr.get_or_add_trPr()
        th = OxmlElement("w:tblHeader")
        th.set(qn("w:val"), "true")
        trPr.append(th)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell_text(t.rows[i].cells[j], val, size=font_size,
                      bold=(first_col_bold and j == 0))
        if zebra and i % 2 == 0:
            for j in range(len(headers)):
                shade_cell(t.rows[i].cells[j], LIGHT_BG)
    if widths_cm:
        for j, w in enumerate(widths_cm):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


def spacer(doc, pts=4):
    add_para(doc, "", size=4, space_after=pts)

def cover_path(bgm_id):
    return os.path.join(COVER_DIR, "%d.jpg" % bgm_id)


def _looks_like_image(path):
    try:
        with open(path, "rb") as f:
            head = f.read(16)
    except Exception:
        return False
    if head.startswith(b"\xff\xd8") or head.startswith(b"\x89PNG"):
        return True
    if head[:6] in (b"GIF87a", b"GIF89a"):
        return True
    return head.startswith(b"RIFF") and head[8:12] == b"WEBP"


def download_cover(bgm_id, url, logger=None, timeout=30):
    """下载封面到本地缓存；失败或已存在则返回路径/None。"""
    if not bgm_id or not url:
        return None
    path = cover_path(bgm_id)
    if os.path.exists(path) and os.path.getsize(path) > 1024 and _looks_like_image(path):
        return path
    try:
        import requests
        r = requests.get(url, headers={"User-Agent": "AnimeAutoBot/1.0"}, timeout=timeout)
        r.raise_for_status()
        ensure_dirs()
        tmp = path + ".tmp"
        with open(tmp, "wb") as f:
            f.write(r.content)
        if _looks_like_image(tmp):
            os.replace(tmp, path)
            return path
        os.remove(tmp)
        if logger:
            logger.warning("封面文件非图片，已丢弃 bgm#%s", bgm_id)
        return None
    except Exception as e:
        if logger:
            logger.warning("封面下载失败 bgm#%s: %s", bgm_id, e)
        return None

def add_hr(doc, color=ACCENT, size=16, space_after=8):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_cover(doc, title, subtitle_lines=(), date_line=None):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    add_para(doc, "", space_after=90)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run(run, size=26, bold=True, color=ACCENT)
    p.paragraph_format.space_after = Pt(8)
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hrPr = hr._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "20")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom)
    hrPr.append(pbdr)
    for line in subtitle_lines:
        add_para(doc, line, size=11, color="666666", align="center", space_after=2)
    if date_line:
        add_para(doc, date_line, size=11, color="888888", align="center", space_after=2)
    add_para(doc, "", space_after=24)
    doc.add_page_break()


def _field_run(p, instr):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    r1 = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r1._r.append(f1)
    r2 = p.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2._r.append(it)
    r3 = p.add_run()
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end"); r3._r.append(f2)


def add_footer_pagenum(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第 "); set_run(run, size=8, color="999999")
    _field_run(p, " PAGE ")
    run = p.add_run(" 页 / 共 "); set_run(run, size=8, color="999999")
    _field_run(p, " NUMPAGES ")
    run = p.add_run(" 页"); set_run(run, size=8, color="999999")


def colorize_score_cell(cell):
    """按分数给单元格着色：>=8.5 绿，>=7.5 橙，其余灰。"""
    from docx.shared import RGBColor
    try:
        v = float((cell.text or "").strip())
    except ValueError:
        return
    color = "1E8E3E" if v >= 8.5 else ("E8930C" if v >= 7.5 else "9A9A9A")
    for para in cell.paragraphs:
        for r in para.runs:
            r.font.color.rgb = RGBColor.from_string(color)