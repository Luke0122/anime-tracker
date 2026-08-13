# -*- coding: utf-8 -*-
"""生成季度评分汇总 Word 文档（含封面与美化排版）。"""
import datetime
import os
import sys

from docx.shared import Cm, Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from common import (add_cover, add_footer_pagenum, add_heading, add_para, colorize_score_cell,
                    download_cover, init_doc, make_table, season_label, set_run, setup_logging,
                    spacer)



def _dedupe_top(rated, limit=10):
    """同一 bangumi 条目只保留一次（优先保留不带 # 的主标题），再按评分排序。"""
    best = {}
    for x in rated:
        bid = x.get("bgm_id")
        key = ("b", bid) if bid else ("n", x.get("title", ""))
        cur = best.get(key)
        if cur is None:
            best[key] = x
            continue
        def tk(t):
            t = t or ""
            return (1 if "#" in t else 0, -len(t))
        if tk(x.get("title", "")) < tk(cur.get("title", "")):
            best[key] = x
    vals = list(best.values())
    vals.sort(key=lambda x: (-(x.get("rating") or 0), -(x.get("rating_total") or 0)))
    return vals[:limit]

def build(bgm_items, year, start_month, phase_label, out_path, generated_at=None):
    """phase_label: 初期/中期/末期/补做"""
    generated_at = generated_at or datetime.date.today().isoformat()
    label = season_label(year, start_month)
    items = list(bgm_items)

    rated = [x for x in items if x.get("rating") is not None]
    unrated = [x for x in items if x.get("rating") is None]
    rated.sort(key=lambda x: (-(x.get("rating") or 0), -(x.get("rating_total") or 0)))

    n9 = sum(1 for x in rated if x["rating"] >= 9.0)
    n8 = sum(1 for x in rated if 8.0 <= x["rating"] < 9.0)
    n7 = sum(1 for x in rated if 7.0 <= x["rating"] < 8.0)
    nlow = sum(1 for x in rated if x["rating"] < 7.0)

    doc = init_doc()
    doc.core_properties.title = "%s评分汇总（%s）" % (label, phase_label)
    doc.core_properties.author = "Anime Auto"
    add_cover(
        doc,
        "%s评分汇总（%s）" % (label, phase_label),
        ["日本 TV 动画 · 本季评分跟踪", "数据来源：Bangumi 评分（bgm.tv）"],
        date_line="生成日期：%s" % generated_at,
    )
    add_footer_pagenum(doc)

    add_heading(doc, "一、评分总表", level=1)
    rows = []
    for i, x in enumerate(rated, start=1):
        rows.append([
            str(i),
            x.get("title", ""),
            "%.1f" % x["rating"],
            str(x.get("rating_total") or 0),
            x.get("weekday", ""),
            x.get("time", ""),
        ])
    for x in unrated:
        rows.append(["—", x.get("title", ""), "暂无评分", "—", x.get("weekday", ""), x.get("time", "")])
    t = make_table(doc,
                   ["排名", "标题", "评分", "评分人数", "星期", "时间"],
                   rows,
                   widths_cm=[1.2, 7.2, 1.8, 2.0, 1.5, 1.6],
                   font_size=8.5)
    for row in t.rows[1:]:
        colorize_score_cell(row.cells[2])
    spacer(doc)

    add_heading(doc, "二、分档统计", level=1)
    stats = [
        ["9.0 及以上", "%d 部" % n9],
        ["8.0 – 8.9", "%d 部" % n8],
        ["7.0 – 7.9", "%d 部" % n7],
        ["7.0 以下", "%d 部" % nlow],
        ["暂无评分", "%d 部" % len(unrated)],
        ["合计", "%d 部" % len(items)],
    ]
    make_table(doc, ["分档", "数量"], stats, widths_cm=[6.0, 9.5], font_size=10)
    spacer(doc)

    add_heading(doc, "三、高分亮点（Top 10）", level=1)
    logger = setup_logging()
    for i, x in enumerate(_dedupe_top(rated, 10), start=1):
        p = doc.add_paragraph()
        cover = download_cover(x.get("bgm_id"), x.get("image_url"), logger=logger)
        if cover:
            try:
                run = p.add_run()
                run.add_picture(cover, width=Cm(1.7))
                p.add_run("  ")
            except Exception as e:
                logger.warning("封面插入失败「%s」：%s", x.get("title"), e)
        run2 = p.add_run("%d. %s　——　%.1f 分（%d 人评分）" %
                         (i, x.get("title", ""), x["rating"], x.get("rating_total") or 0))
        set_run(run2, size=10, bold=(i <= 3))
        p.paragraph_format.space_after = Pt(6)

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path